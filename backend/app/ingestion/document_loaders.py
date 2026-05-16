import csv
import re
import warnings as py_warnings
from dataclasses import dataclass, field
from pathlib import Path
from typing import Dict, List, Optional
from uuid import uuid4

from backend.app.ingestion.chunking import ChunkedText, chunk_text_with_metadata, summarize
from backend.app.ingestion.contextualizer import Contextualizer
from backend.app.ingestion.ocr_client import OcrClient
from backend.app.ingestion.section_refiner import SectionRefiner
from backend.app.pipeline.schemas import (
    ReferenceDocument,
    ReferenceItem,
    SourceChunk,
    SourceDocument,
    TemplateBlock,
    TemplateSection,
    TemplateSectionCandidate,
    TemplateRefinementSuggestion,
    TemplateStructure,
)


HEADING_PATTERNS = [
    re.compile(r"^\s*(第[一二三四五六七八九十\d]+[章節])\s*(.+)$"),
    re.compile(r"^\s*([一二三四五六七八九十]+、)\s*(.+)$"),
    re.compile(r"^\s*(\d+(?:\.\d+)*)[.)、]?\s+(.+)$"),
]

SOP_HEADING_KEYWORDS = [
    "purpose",
    "scope",
    "definition",
    "responsibility",
    "safety",
    "tools",
    "materials",
    "procedure",
    "troubleshooting",
    "diagnosis",
    "repair",
    "maintenance",
    "inspection",
    "acceptance",
    "record",
    "revision",
    "目的",
    "範圍",
    "范围",
    "適用",
    "适用",
    "定義",
    "定义",
    "職責",
    "职责",
    "安全",
    "工具",
    "材料",
    "程序",
    "流程",
    "步驟",
    "步骤",
    "維修",
    "维修",
    "檢修",
    "检修",
    "故障",
    "診斷",
    "诊断",
    "驗收",
    "验收",
    "記錄",
    "记录",
    "文件",
    "注意事項",
    "注意事项",
]

METADATA_LABEL_PATTERNS = [
    re.compile(pattern, re.IGNORECASE)
    for pattern in [
        r"^(document\s*)?title$",
        r"^(document\s*)?(no\.?|number|id)$",
        r"^(version|revision|rev\.?|effective\s*date|date)$",
        r"^(prepared|reviewed|approved)(\s*by)?$",
        r"^(author|owner|department|page)$",
        r"^(文件)?(標題|标题|名稱|名称|編號|编号)$",
        r"^(版本|版次|修訂|修订|日期|生效日期)$",
        r"^(製作|制作者|製表|审核|審核|核准|批准)(人|者)?$",
    ]
]

DOCUMENT_TITLE_TERMS = [
    "standard operating procedure",
    "sop",
    "work instruction",
    "維修作業標準",
    "维修作业标准",
    "標準作業程序",
    "标准作业程序",
    "作業指導書",
    "作业指导书",
]


@dataclass
class DocxTextBlock:
    text: str
    style_name: str = ""
    source: str = "paragraph"
    metadata: Dict[str, str] = field(default_factory=dict)


def load_source_pdf(
    path: str,
    ocr_client: OcrClient = None,
    chunk_size: int = 900,
    chunk_overlap: int = 120,
    chunk_method: str = "vanilla",
    contextualizer: Optional[Contextualizer] = None,
) -> SourceDocument:
    file_path = Path(path)
    warnings: List[str] = []
    raw_text, extraction_method = _extract_pdf_text(file_path, warnings, ocr_client)
    if not raw_text.strip():
        warnings.append("No source text extracted.")
    document_id = f"source-{uuid4().hex[:10]}"
    chunked = chunk_text_with_metadata(raw_text, chunk_size=chunk_size, overlap=chunk_overlap)
    chunks = []
    document_context = _document_context(raw_text, chunk_method, contextualizer)
    for index, chunk in enumerate(chunked):
        embedding_text = _embedding_text(raw_text, chunk, chunk_method, contextualizer, document_context)
        metadata = dict(chunk.metadata)
        metadata.update({"chunk_method": chunk_method, "extraction_method": extraction_method})
        chunks.append(
            SourceChunk(
                chunk_id=f"{document_id}-chunk-{index}",
                document_id=document_id,
                file_name=file_path.name,
                content=chunk.text,
                embedding_text=embedding_text,
                summary=summarize(chunk.text),
                page_start=chunk.page_start,
                page_end=chunk.page_end,
                metadata=metadata,
            )
        )
    return SourceDocument(
        document_id=document_id,
        file_name=file_path.name,
        raw_text=raw_text,
        chunks=chunks,
        metadata={"extraction_method": extraction_method},
        warnings=warnings,
    )


def load_reference_file(
    path: str,
    ocr_client: OcrClient = None,
    chunk_size: int = 900,
    chunk_overlap: int = 120,
    chunk_method: str = "vanilla",
    contextualizer: Optional[Contextualizer] = None,
) -> ReferenceDocument:
    file_path = Path(path)
    ext = file_path.suffix.lower()
    if ext in {".xlsx", ".xls"}:
        return _load_excel(file_path)
    if ext == ".pdf":
        source_like = load_source_pdf(
            path,
            ocr_client=ocr_client,
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            chunk_method=chunk_method,
            contextualizer=contextualizer,
        )
        return ReferenceDocument(
            document_id=source_like.document_id.replace("source-", "ref-"),
            file_name=file_path.name,
            file_type="pdf",
            items=[
                ReferenceItem(
                    item_id=chunk.chunk_id.replace("source-", "ref-"),
                    document_id=source_like.document_id.replace("source-", "ref-"),
                    file_name=file_path.name,
                    item_type="unstructured_chunk",
                    content=chunk.content,
                    embedding_text=chunk.embedding_text,
                    summary=chunk.summary,
                    location=_page_location(chunk.page_start, chunk.page_end),
                    metadata=chunk.metadata,
                )
                for chunk in source_like.chunks
            ],
            warnings=source_like.warnings,
        )
    if ext == ".md":
        return _load_text_reference(file_path, "markdown", chunk_size, chunk_overlap, chunk_method, contextualizer)
    return _load_text_reference(file_path, "txt", chunk_size, chunk_overlap, chunk_method, contextualizer)


def load_template_docx(
    path: str,
    section_detection_mode: str = "rules",
    section_refiner: Optional[SectionRefiner] = None,
) -> TemplateStructure:
    file_path = Path(path)
    warnings: List[str] = []
    sections: List[TemplateSection] = []
    blocks: List[TemplateBlock] = []
    candidates: List[TemplateSectionCandidate] = []
    template_id = f"template-{uuid4().hex[:10]}"
    try:
        from docx import Document

        document = Document(str(file_path))
        docx_blocks = _docx_text_blocks(document)
        blocks = [
            TemplateBlock(
                block_id=f"block-{index + 1}",
                text=block.text,
                style_name=block.style_name,
                source_type=block.source,
                order_index=index,
                metadata=block.metadata,
            )
            for index, block in enumerate(docx_blocks)
        ]
        table_block_count = sum(1 for block in blocks if block.source_type == "table")
        if table_block_count:
            warnings.append(f"Read {table_block_count} text block(s) from DOCX tables for section detection.")
        candidates = _section_candidates_from_blocks(blocks)
        sections = _sections_from_candidates(candidates, blocks)
    except Exception as exc:
        warnings.append(f"DOCX heading extraction failed: {exc}")
        try:
            text = file_path.read_text(encoding="utf-8", errors="replace")
        except Exception:
            text = ""
        blocks = [
            TemplateBlock(
                block_id=f"line-{index + 1}",
                text=line.strip(),
                source_type="text",
                order_index=index,
            )
            for index, line in enumerate(text.splitlines())
            if line.strip()
        ]
        candidates = _section_candidates_from_blocks(blocks)
        sections = _sections_from_candidates(candidates, blocks)

    if not sections:
        warnings.append("No clear headings detected; using a single General SOP section.")
        sections = [
            TemplateSection(
                section_id="section-1",
                title="General SOP",
                level=1,
                start_block_index=0,
                existing_text="",
            )
        ]
    suggestions: List[TemplateRefinementSuggestion] = []
    refinement_mode = "rules"
    feedback_intent = "guidance"
    if section_detection_mode == "rules_llm":
        if section_refiner:
            refined = section_refiner.refine(
                file_name=file_path.name,
                template_id=template_id,
                blocks=blocks,
                candidates=candidates,
                fallback_sections=sections,
            )
            sections = refined.sections
            suggestions = refined.refinement_suggestions
            warnings.extend(refined.warnings)
            refinement_mode = refined.refinement_mode
            feedback_intent = refined.feedback_intent
        if refinement_mode == "llm":
            warnings.append("LLM section refinement produced the review proposal.")
        else:
            warnings.append("LLM section refinement unavailable; rule-based sections are shown.")

    return TemplateStructure(
        template_id=template_id,
        file_name=file_path.name,
        sections=sections,
        warnings=warnings,
        refinement_suggestions=suggestions,
        blocks=blocks,
        candidates=candidates,
        resolution_id=f"resolution-{uuid4().hex[:10]}",
        refinement_mode=refinement_mode,
        feedback_intent=feedback_intent,
    )


def _load_text_reference(
    file_path: Path,
    file_type: str,
    chunk_size: int,
    chunk_overlap: int,
    chunk_method: str,
    contextualizer: Optional[Contextualizer],
) -> ReferenceDocument:
    warnings: List[str] = []
    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
    except Exception as exc:
        warnings.append(f"Could not read text file: {exc}")
        text = ""
    document_id = f"ref-{uuid4().hex[:10]}"
    items = []
    document_context = _document_context(text, chunk_method, contextualizer)
    for index, chunk in enumerate(chunk_text_with_metadata(text, chunk_size=chunk_size, overlap=chunk_overlap)):
        embedding_text = _embedding_text(text, chunk, chunk_method, contextualizer, document_context)
        metadata = dict(chunk.metadata)
        metadata["chunk_method"] = chunk_method
        items.append(
            ReferenceItem(
                item_id=f"{document_id}-item-{index}",
                document_id=document_id,
                file_name=file_path.name,
                item_type="unstructured_chunk",
                content=chunk.text,
                embedding_text=embedding_text,
                summary=summarize(chunk.text),
                location=f"chunk {index + 1}",
                metadata=metadata,
            )
        )
    return ReferenceDocument(
        document_id=document_id,
        file_name=file_path.name,
        file_type=file_type,
        items=items,
        warnings=warnings,
    )


def _load_excel(file_path: Path) -> ReferenceDocument:
    warnings: List[str] = []
    document_id = f"ref-{uuid4().hex[:10]}"
    rows: List[Dict[str, str]] = []
    try:
        import openpyxl

        with py_warnings.catch_warnings():
            py_warnings.filterwarnings(
                "ignore",
                message="Data Validation extension is not supported and will be removed",
                category=UserWarning,
                module="openpyxl.worksheet._reader",
            )
            workbook = openpyxl.load_workbook(str(file_path), data_only=True)
        sheet = workbook.active
        headers = [str(cell.value or "").strip() for cell in next(sheet.iter_rows(min_row=1, max_row=1))]
        for row in sheet.iter_rows(min_row=2, values_only=True):
            rows.append({headers[i] or f"column_{i + 1}": "" if value is None else str(value) for i, value in enumerate(row)})
    except Exception as exc:
        warnings.append(f"openpyxl failed, trying csv-style read: {exc}")
        try:
            with file_path.open("r", encoding="utf-8", errors="replace") as handle:
                rows = list(csv.DictReader(handle))
        except Exception as csv_exc:
            warnings.append(f"Could not parse Excel/reference table: {csv_exc}")
    items: List[ReferenceItem] = []
    for index, row in enumerate(rows):
        raw = {k: v for k, v in row.items() if str(v).strip()}
        content = "\n".join(f"{key}: {value}" for key, value in raw.items())
        items.append(
            ReferenceItem(
                item_id=f"{document_id}-row-{index + 1}",
                document_id=document_id,
                file_name=file_path.name,
                item_type="structured_record",
                content=content,
                summary=summarize(content),
                location=f"row {index + 2}",
                metadata={"raw_fields": str(raw), "field_count": str(len(raw))},
            )
        )
    return ReferenceDocument(
        document_id=document_id,
        file_name=file_path.name,
        file_type="excel",
        items=items,
        warnings=warnings,
    )


def _heading_level(style_name: str, text: str) -> int:
    normalized = style_name.lower()
    if _is_metadata_label(text):
        return 0
    if "heading 1" in normalized or "標題 1" in style_name:
        return 1
    if "heading 2" in normalized or "標題 2" in style_name:
        return 2
    for pattern in HEADING_PATTERNS:
        match = pattern.match(text)
        if match:
            marker = match.group(1)
            return 2 if "." in marker else 1
    lowered = text.lower().strip()
    if (
        1 <= len(text) <= 80
        and not re.search(r"[。．.;；]", text)
        and len(text.split()) <= 6
        and any(keyword in lowered or keyword in text for keyword in SOP_HEADING_KEYWORDS)
    ):
        return 1
    return 0


def _section_candidates_from_blocks(blocks: List[TemplateBlock]) -> List[TemplateSectionCandidate]:
    candidates: List[TemplateSectionCandidate] = []
    for block in blocks:
        if _should_ignore_as_section(block):
            continue
        level = _heading_level(block.style_name or "", block.text)
        if not level:
            continue
        detector = "style" if block.style_name and "heading" in block.style_name.lower() else "rules"
        warnings = _candidate_warnings(block)
        confidence = 0.9 if detector == "style" else 0.72
        if warnings:
            confidence = min(confidence, 0.45)
        candidates.append(
            TemplateSectionCandidate(
                candidate_id=f"candidate-{len(candidates) + 1}",
                title=block.text.strip(),
                level=level,
                source_block_ids=[block.block_id],
                confidence=confidence,
                reason=f"{detector} matched this block as a likely SOP section heading.",
                detector=detector,
                warnings=warnings,
                metadata=dict(block.metadata),
            )
        )
    return candidates


def _sections_from_candidates(
    candidates: List[TemplateSectionCandidate], blocks: List[TemplateBlock]
) -> List[TemplateSection]:
    block_index_by_id = {block.block_id: index for index, block in enumerate(blocks)}
    sections: List[TemplateSection] = []
    for index, candidate in enumerate(candidates):
        start_index = block_index_by_id.get(candidate.source_block_ids[0], index) if candidate.source_block_ids else index
        next_candidate = candidates[index + 1] if index + 1 < len(candidates) else None
        next_start = (
            block_index_by_id.get(next_candidate.source_block_ids[0], len(blocks))
            if next_candidate and next_candidate.source_block_ids
            else len(blocks)
        )
        existing_text = "\n".join(
            block.text.strip()
            for block in blocks[start_index + 1 : next_start]
            if block.text.strip()
        )
        sections.append(
            TemplateSection(
                section_id=f"section-{len(sections) + 1}",
                title=candidate.title,
                level=candidate.level,
                start_block_index=start_index,
                end_block_index=max(next_start - 1, start_index),
                existing_text=existing_text,
                style_name=blocks[start_index].style_name if start_index < len(blocks) else None,
                source_block_ids=candidate.source_block_ids,
                confidence=candidate.confidence,
                operation="keep",
                reason=candidate.reason,
            )
        )
    return sections


def _docx_text_blocks(document) -> List[DocxTextBlock]:
    blocks: List[DocxTextBlock] = []
    for item in _iter_docx_block_items(document):
        if hasattr(item, "text"):
            text = item.text.strip()
            if not text:
                continue
            style_name = item.style.name if item.style else ""
            blocks.append(DocxTextBlock(text=text, style_name=style_name, source="paragraph"))
            continue

        blocks.extend(_table_text_blocks(item))
    return _annotate_docx_blocks(blocks)


def _iter_docx_block_items(parent):
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    if hasattr(parent, "element") and hasattr(parent.element, "body"):
        parent_element = parent.element.body
    elif hasattr(parent, "_tc"):
        parent_element = parent._tc
    else:
        parent_element = parent.element
    for child in parent_element.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent)
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent)


def _table_text_blocks(table) -> List[DocxTextBlock]:
    blocks: List[DocxTextBlock] = []
    seen_cells = set()
    for row_index, row in enumerate(table.rows, start=1):
        for column_index, cell in enumerate(row.cells, start=1):
            cell_key = cell._tc
            if cell_key in seen_cells:
                continue
            seen_cells.add(cell_key)
            for item in _iter_docx_block_items(cell):
                if hasattr(item, "text"):
                    text = item.text.strip()
                    if not text:
                        continue
                    style_name = item.style.name if item.style else ""
                    blocks.append(
                        DocxTextBlock(
                            text=text,
                            style_name=style_name,
                            source="table",
                            metadata={"table_row": str(row_index), "table_column": str(column_index)},
                        )
                    )
                    continue
                blocks.extend(_table_text_blocks(item))
    return blocks


def _annotate_docx_blocks(blocks: List[DocxTextBlock]) -> List[DocxTextBlock]:
    total = max(len(blocks), 1)
    for index, block in enumerate(blocks):
        metadata = dict(block.metadata)
        metadata["position_hint"] = _position_hint(index, total)
        if _is_metadata_label(block.text):
            metadata["looks_like_metadata_label"] = "true"
            metadata["section_recommendation"] = "ignore"
        if _looks_like_document_title(block.text, block.style_name, index):
            metadata["looks_like_document_title"] = "true"
            metadata["section_recommendation"] = "ignore"
        block.metadata = metadata
    return blocks


def _position_hint(index: int, total: int) -> str:
    if index < 5 or index / total <= 0.15:
        return "early"
    if index / total >= 0.85:
        return "late"
    return "middle"


def _is_metadata_label(text: str) -> bool:
    normalized = re.sub(r"[:：\s]+", " ", text.strip()).strip()
    if not normalized or len(normalized) > 40:
        return False
    return any(pattern.match(normalized) for pattern in METADATA_LABEL_PATTERNS)


def _looks_like_document_title(text: str, style_name: str, order_index: int) -> bool:
    if order_index > 4:
        return False
    stripped = text.strip()
    if not stripped or _heading_numbered(stripped):
        return False
    lowered = stripped.lower()
    style_lower = (style_name or "").lower()
    if "title" in style_lower and not any(keyword in lowered or keyword in stripped for keyword in SOP_HEADING_KEYWORDS):
        return True
    if any(term in lowered or term in stripped for term in DOCUMENT_TITLE_TERMS):
        return True
    return False


def _heading_numbered(text: str) -> bool:
    return any(pattern.match(text) for pattern in HEADING_PATTERNS)


def _should_ignore_as_section(block: TemplateBlock) -> bool:
    return block.metadata.get("section_recommendation") == "ignore"


def _candidate_warnings(block: TemplateBlock) -> List[str]:
    warnings = []
    if block.metadata.get("looks_like_document_title") == "true":
        warnings.append("This block looks like a document title or cover-page label; verify before using as a fillable section.")
    if block.metadata.get("looks_like_metadata_label") == "true":
        warnings.append("This block looks like an administrative metadata label; verify before using as a fillable section.")
    return warnings


def _sections_from_text(text: str) -> List[TemplateSection]:
    sections = []
    for index, line in enumerate(text.splitlines()):
        if _heading_level("", line.strip()):
            sections.append(
                TemplateSection(
                    section_id=f"section-{len(sections) + 1}",
                    title=line.strip(),
                    level=_heading_level("", line.strip()),
                    start_block_index=index,
                )
            )
    return sections


def _embedding_text(
    document_text: str,
    chunk: ChunkedText,
    chunk_method: str,
    contextualizer: Optional[Contextualizer],
    document_context: str = "",
) -> str:
    if chunk_method == "contextual":
        return f"Document Context: {document_context}\n\nChunk: {chunk.text}" if document_context else chunk.text
    if chunk_method == "anthropic":
        context = contextualizer.chunk_context(document_text, chunk.text) if contextualizer else ""
        return f"{context}\n{chunk.text}" if context else chunk.text
    return chunk.text


def _document_context(document_text: str, chunk_method: str, contextualizer: Optional[Contextualizer]) -> str:
    if chunk_method != "contextual":
        return ""
    return contextualizer.document_summary(document_text) if contextualizer else summarize(document_text, 500)


def _page_location(page_start: Optional[int], page_end: Optional[int]) -> Optional[str]:
    if not page_start:
        return None
    if page_end and page_end != page_start:
        return f"pages {page_start}-{page_end}"
    return f"page {page_start}"


def _extract_pdf_text(file_path: Path, warnings: List[str], ocr_client: OcrClient = None):
    if ocr_client and ocr_client.is_configured():
        result = ocr_client.extract_pdf_text(str(file_path))
        warnings.extend(result.warnings)
        if result.text.strip():
            return result.text, result.method
        warnings.append("Falling back to PyMuPDF text extraction after OCR produced no usable text.")

    try:
        text = _extract_text_pymupdf(file_path)
        if text.strip():
            return text, "pymupdf"
        warnings.append("PyMuPDF produced no usable text.")
    except Exception as exc:
        warnings.append(f"PyMuPDF text extraction unavailable or failed: {exc}")

    try:
        text = file_path.read_text(encoding="utf-8", errors="replace")
        if text.strip():
            warnings.append("Fell back to plain text read for PDF path.")
            return text, "text_read"
    except Exception as exc:
        warnings.append(f"Plain text fallback failed: {exc}")

    return "", "failed"


def _extract_text_pymupdf(file_path: Path) -> str:
    import fitz

    doc = fitz.open(str(file_path))
    page_texts = []
    try:
        for index, page in enumerate(doc):
            page_text = page.get_text().strip()
            if page_text:
                page_texts.append(f"<!-- Page {index + 1} -->\n{page_text}")
    finally:
        doc.close()
    return "\n\n".join(page_texts)
