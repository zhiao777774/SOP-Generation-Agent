from pathlib import Path
from typing import List, Optional

from backend.app.pipeline.schemas import (
    GenerationResult,
    StructuredBlock,
    StructuredListItem,
    TemplateBlock,
    TemplateStructure,
)


class DocxRenderer:
    def __init__(self):
        try:
            from markdown_it import MarkdownIt

            self.markdown = MarkdownIt("commonmark", options_update={"html": False})
        except Exception:
            self.markdown = None

    def render(
        self,
        template_path: Optional[str],
        template_structure: TemplateStructure,
        generation: GenerationResult,
        output_path: Path,
    ) -> Path:
        try:
            from docx import Document
        except Exception as exc:
            output_path.write_text(
                f"DOCX dependency unavailable: {exc}\n\n{self._plain_text(generation)}",
                encoding="utf-8",
            )
            return output_path

        if template_path and Path(template_path).suffix.lower() == ".docx" and Path(template_path).exists():
            document = Document(template_path)
        else:
            document = Document()
            document.add_heading(template_structure.file_name or "Generated SOP", level=1)

        section_by_id = {section.section_id: section for section in generation.sections}
        inserted = set()
        anchors_by_block_id = self._anchors_by_template_block_id(document, template_structure.blocks)
        for template_section in template_structure.sections:
            draft = section_by_id.get(template_section.section_id)
            if not draft:
                continue
            anchor = self._anchor_for_section(template_section.source_block_ids, anchors_by_block_id)
            if anchor is not None:
                self._insert_blocks_after(anchor, draft.blocks)
                inserted.add(draft.section_id)

        for draft in generation.sections:
            if draft.section_id in inserted:
                continue
            document.add_heading(draft.title, level=1)
            for block in draft.blocks:
                self._append_block(document, block)

        output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(output_path))
        return output_path

    def _insert_blocks_after(self, anchor, blocks: List[StructuredBlock]) -> None:
        current = anchor
        for block in blocks:
            current = self._insert_block_after(current, block)

    def _insert_block_after(self, current, block: StructuredBlock):
        parent = current._parent
        if block.block_type == "table":
            table = self._add_table(parent, block)
            self._element(current).addnext(table._tbl)
            return table
        if block.block_type == "image":
            return self._insert_image_after(current, block)

        inserted = []
        for paragraph_spec in self._paragraph_specs(block):
            paragraph = self._insert_paragraph_after(current)
            style_applied = self._apply_style(paragraph, paragraph_spec["style"])
            self._apply_indent(paragraph, paragraph_spec["indent"])
            self._write_inline(paragraph, self._fallback_content(paragraph_spec, style_applied))
            inserted.append(paragraph)
            current = paragraph
        return inserted[-1] if inserted else current

    def _insert_paragraph_after(self, current, style: Optional[str] = None):
        from docx.oxml import OxmlElement
        from docx.text.paragraph import Paragraph

        new_p = OxmlElement("w:p")
        self._element(current).addnext(new_p)
        new_paragraph = Paragraph(new_p, current._parent)
        self._apply_style(new_paragraph, style)
        return new_paragraph

    def _append_block(self, document, block: StructuredBlock):
        if block.block_type == "table":
            return self._add_table(document, block)
        if block.block_type == "image":
            return self._append_image(document, block)

        current = None
        for paragraph_spec in self._paragraph_specs(block):
            paragraph = document.add_paragraph()
            style_applied = self._apply_style(paragraph, paragraph_spec["style"])
            self._apply_indent(paragraph, paragraph_spec["indent"])
            self._write_inline(paragraph, self._fallback_content(paragraph_spec, style_applied))
            current = paragraph
        return current

    def _insert_image_after(self, current, block: StructuredBlock):
        image_paragraph = self._insert_paragraph_after(current)
        self._write_image_or_fallback(image_paragraph, block)
        caption = self._content(block) or block.caption_md
        if caption:
            caption_paragraph = self._insert_paragraph_after(image_paragraph)
            self._write_inline(caption_paragraph, caption)
            return caption_paragraph
        return image_paragraph

    def _append_image(self, document, block: StructuredBlock):
        image_paragraph = document.add_paragraph()
        self._write_image_or_fallback(image_paragraph, block)
        caption = self._content(block) or block.caption_md
        if caption:
            caption_paragraph = document.add_paragraph()
            self._write_inline(caption_paragraph, caption)
            return caption_paragraph
        return image_paragraph

    def _write_image_or_fallback(self, paragraph, block: StructuredBlock) -> None:
        image_path = Path(block.image_path) if block.image_path else None
        if not image_path or not image_path.exists():
            self._write_inline(paragraph, block.alt_text or block.caption_md or "[image unavailable]")
            return
        try:
            from docx.shared import Inches

            paragraph.add_run().add_picture(str(image_path), width=Inches(5.8))
        except Exception:
            self._write_inline(paragraph, block.alt_text or block.caption_md or "[image unavailable]")

    def _paragraph_specs(self, block: StructuredBlock) -> List[dict]:
        block_type = block.block_type
        if block_type == "heading":
            return [
                {
                    "style": f"Heading {self._heading_level(block.level)}",
                    "content": self._content(block),
                    "indent": 0,
                }
            ]
        if block_type in {"bullet", "bullet_list"}:
            return self._list_specs(block.items, "List Bullet", fallback=self._content(block))
        if block_type in {"numbered", "numbered_list"}:
            return self._list_specs(block.items, "List Number", fallback=self._content(block))
        if block_type == "callout":
            prefix = block.callout_type.upper() if block.callout_type else "NOTE"
            return [{"style": None, "content": f"{prefix}: {self._content(block)}", "indent": 0}]
        if block_type == "image":
            return []
        return [{"style": None, "content": self._content(block), "indent": 0}]

    def _list_specs(
        self,
        items: List[StructuredListItem],
        style: str,
        fallback: str = "",
        depth: int = 0,
    ) -> List[dict]:
        if not items and fallback:
            return [{"style": style, "content": fallback, "indent": depth, "fallback_prefix": self._list_prefix(style, 1)}]
        specs = []
        for index, item in enumerate(items, start=1):
            specs.append(
                {
                    "style": style,
                    "content": self._item_content(item),
                    "indent": depth,
                    "fallback_prefix": self._list_prefix(style, index),
                }
            )
            specs.extend(self._list_specs(item.items, style, depth=depth + 1))
        return specs

    def _add_table(self, parent, block: StructuredBlock):
        rows = block.rows or []
        headers = block.headers or []
        column_count = max([len(headers)] + [len(row) for row in rows] + [1])
        table = self._new_table(parent, len(rows) + (1 if headers else 0), column_count)
        try:
            table.style = "Table Grid"
        except Exception:
            pass
        row_offset = 0
        if headers:
            self._fill_table_row(table.rows[0], headers, bold=True)
            row_offset = 1
        for index, row in enumerate(rows):
            self._fill_table_row(table.rows[index + row_offset], row, bold=False)
        return table

    def _fill_table_row(self, row, values: List[str], bold: bool) -> None:
        for index, cell in enumerate(row.cells):
            value = values[index] if index < len(values) else ""
            paragraph = cell.paragraphs[0]
            paragraph.clear()
            self._write_inline(paragraph, value, force_bold=bold)

    def _write_inline(self, paragraph, content: str, force_bold: bool = False) -> None:
        if not self.markdown:
            run = paragraph.add_run(content)
            run.bold = force_bold or None
            return

        tokens = self.markdown.parseInline(content, {})
        children = tokens[0].children if tokens and tokens[0].children else []
        strong = 0
        emphasis = 0
        for token in children:
            if token.type == "strong_open":
                strong += 1
            elif token.type == "strong_close":
                strong = max(strong - 1, 0)
            elif token.type == "em_open":
                emphasis += 1
            elif token.type == "em_close":
                emphasis = max(emphasis - 1, 0)
            elif token.type in {"text", "code_inline"}:
                run = paragraph.add_run(token.content)
                run.bold = force_bold or bool(strong)
                run.italic = bool(emphasis)
                if token.type == "code_inline":
                    run.font.name = "Consolas"
            elif token.type in {"softbreak", "hardbreak"}:
                paragraph.add_run().add_break()

    def _apply_indent(self, paragraph, depth: int) -> None:
        if depth <= 0:
            return
        try:
            from docx.shared import Inches

            paragraph.paragraph_format.left_indent = Inches(0.25 * depth)
        except Exception:
            return

    def _new_table(self, parent, rows: int, cols: int):
        try:
            return parent.add_table(rows=rows, cols=cols)
        except TypeError:
            from docx.shared import Inches

            return parent.add_table(rows=rows, cols=cols, width=Inches(6))

    def _apply_style(self, paragraph, style: Optional[str]) -> bool:
        if not style:
            return True
        try:
            paragraph.style = style
            return True
        except Exception:
            return False

    def _fallback_content(self, paragraph_spec: dict, style_applied: bool) -> str:
        content = paragraph_spec["content"]
        if style_applied:
            return content
        prefix = paragraph_spec.get("fallback_prefix", "")
        return f"{prefix}{content}" if prefix else content

    def _list_prefix(self, style: str, index: int) -> str:
        if style == "List Number":
            return f"{index}. "
        return "- "

    def _anchors_by_template_block_id(self, document, blocks: List[TemplateBlock]) -> dict:
        anchors = {}
        block_index = 0
        for item in self._iter_docx_block_items(document):
            if self._is_paragraph_like(item):
                block_index = self._map_paragraph_anchor(item, blocks, block_index, anchors)
                continue
            block_index = self._skip_table_blocks(item, blocks, block_index)
        return anchors

    def _map_paragraph_anchor(self, paragraph, blocks: List[TemplateBlock], block_index: int, anchors: dict) -> int:
        text = paragraph.text.strip()
        if not text:
            return block_index
        block = self._matching_block(blocks, block_index, text, "paragraph")
        if block:
            anchors[block.block_id] = paragraph
        return block_index + 1

    def _skip_table_blocks(self, table, blocks: List[TemplateBlock], block_index: int) -> int:
        seen_cells = set()
        for row in table.rows:
            for cell in row.cells:
                cell_key = cell._tc
                if cell_key in seen_cells:
                    continue
                seen_cells.add(cell_key)
                for item in self._iter_docx_block_items(cell):
                    if self._is_paragraph_like(item):
                        text = item.text.strip()
                        if text:
                            self._matching_block(blocks, block_index, text, "table")
                            block_index += 1
                        continue
                    block_index = self._skip_table_blocks(item, blocks, block_index)
        return block_index

    def _matching_block(
        self,
        blocks: List[TemplateBlock],
        start_index: int,
        text: str,
        source_type: str,
    ) -> Optional[TemplateBlock]:
        for index in range(start_index, min(start_index + 8, len(blocks))):
            block = blocks[index]
            if block.source_type == source_type and block.text.strip() == text:
                return block
        return None

    def _anchor_for_section(self, block_ids: List[str], anchors_by_block_id: dict):
        for block_id in block_ids:
            anchor = anchors_by_block_id.get(block_id)
            if anchor is not None:
                return anchor
        return None

    def _iter_docx_block_items(self, parent):
        from docx.oxml.table import CT_Tbl
        from docx.oxml.text.paragraph import CT_P
        from docx.table import Table
        from docx.text.paragraph import Paragraph

        if hasattr(parent, "element") and hasattr(parent.element, "body"):
            parent_element = parent.element.body
        elif hasattr(parent, "_tc"):
            parent_element = parent._tc
        else:
            parent_element = parent.element
        for child in parent_element.iterchildren():
            if isinstance(child, CT_P):
                yield Paragraph(child, parent)
            elif isinstance(child, CT_Tbl):
                yield Table(child, parent)

    def _is_paragraph_like(self, item) -> bool:
        return hasattr(item, "text")

    def _content(self, block: StructuredBlock) -> str:
        return (block.content_md or block.text or "").strip()

    def _item_content(self, item: StructuredListItem) -> str:
        return (item.content_md or item.text or "").strip()

    def _heading_level(self, level: int) -> int:
        if level <= 0:
            return 2
        return min(max(level, 1), 6)

    def _element(self, item):
        paragraph_element = getattr(item, "_p", None)
        if paragraph_element is not None:
            return paragraph_element
        table_element = getattr(item, "_tbl", None)
        if table_element is not None:
            return table_element
        return item._element

    def _plain_text(self, generation: GenerationResult) -> str:
        lines = []
        for section in generation.sections:
            lines.append(section.title)
            for block in section.blocks:
                lines.extend(self._plain_block_lines(block))
            lines.append("")
        return "\n".join(lines)

    def _plain_block_lines(self, block: StructuredBlock) -> List[str]:
        if block.block_type in {"bullet", "bullet_list", "numbered", "numbered_list"}:
            return [self._item_content(item) for item in block.items] or [self._content(block)]
        if block.block_type == "table":
            return ["\t".join(row) for row in ([block.headers] if block.headers else []) + block.rows]
        if block.block_type == "image":
            return [block.caption_md or block.alt_text or "[image]"]
        return [self._content(block)]
