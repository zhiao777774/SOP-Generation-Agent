from pathlib import Path

from docx import Document

from backend.app.pipeline.schemas import (
    GenerationResult,
    StructuredBlock,
    StructuredListItem,
    StructuredSectionDraft,
    TemplateSection,
    TemplateStructure,
)
from backend.app.rendering.docx_renderer import DocxRenderer
from backend.app.reports.report_builders import build_provenance_report


def render_docx(tmp_path: Path, blocks: list[StructuredBlock]) -> Path:
    output_path = tmp_path / "output.docx"
    DocxRenderer().render(
        None,
        TemplateStructure(
            template_id="template-1",
            file_name="template.docx",
            sections=[TemplateSection(section_id="s1", title="Repair Procedure", level=1)],
        ),
        GenerationResult(
            job_id="job-1",
            sections=[
                StructuredSectionDraft(
                    section_id="s1",
                    title="Repair Procedure",
                    blocks=blocks,
                )
            ],
        ),
        output_path,
    )
    return output_path


def test_docx_renderer_converts_inline_markdown_to_runs(tmp_path):
    path = render_docx(
        tmp_path,
        [
            StructuredBlock(
                block_id="b1",
                block_type="paragraph",
                content_md="Inspect **main valve** and record `ERR-207`.",
                source_chunk_ids=["source-1-c1"],
            )
        ],
    )

    document = Document(path)
    paragraph = next(p for p in document.paragraphs if "Inspect" in p.text)

    assert paragraph.text == "Inspect main valve and record ERR-207."
    assert "**" not in paragraph.text
    assert any(run.text == "main valve" and run.bold for run in paragraph.runs)
    assert any(run.text == "ERR-207" and run.font.name for run in paragraph.runs)


def test_docx_renderer_writes_headings_lists_and_tables(tmp_path):
    path = render_docx(
        tmp_path,
        [
            StructuredBlock(block_id="h1", block_type="heading", content_md="Safety checks", level=3),
            StructuredBlock(
                block_id="l1",
                block_type="bullet_list",
                items=[
                    StructuredListItem(content_md="Wear **ESD gloves**.", source_chunk_ids=["source-1-c1"]),
                    StructuredListItem(
                        content_md="Inspect feeder.",
                        items=[StructuredListItem(content_md="Confirm sensor is clean.")],
                    ),
                ],
            ),
            StructuredBlock(
                block_id="n1",
                block_type="numbered_list",
                items=[
                    StructuredListItem(content_md="Power off machine."),
                    StructuredListItem(content_md="Release residual pressure."),
                ],
            ),
            StructuredBlock(
                block_id="t1",
                block_type="table",
                headers=["Item", "Action"],
                rows=[["Valve", "**Replace** if leakage is found"], ["Sensor", "Clean AOI window"]],
            ),
        ],
    )

    document = Document(path)
    styles = [paragraph.style.name for paragraph in document.paragraphs if paragraph.text]

    assert "Heading 3" in styles
    assert any(style == "List Bullet" for style in styles)
    assert any(style == "List Number" for style in styles)
    assert document.tables[0].cell(0, 0).text == "Item"
    assert document.tables[0].cell(1, 1).text == "Replace if leakage is found"
    assert "**" not in document.tables[0].cell(1, 1).text


def test_docx_renderer_falls_back_when_list_style_is_missing(tmp_path, monkeypatch):
    renderer = DocxRenderer()
    original_apply_style = renderer._apply_style

    def fake_apply_style(paragraph, style):
        if style in {"List Bullet", "List Number"}:
            return False
        return original_apply_style(paragraph, style)

    monkeypatch.setattr(renderer, "_apply_style", fake_apply_style)
    output_path = tmp_path / "output.docx"
    renderer.render(
        None,
        TemplateStructure(
            template_id="template-1",
            file_name="template.docx",
            sections=[TemplateSection(section_id="s1", title="Repair Procedure", level=1)],
        ),
        GenerationResult(
            job_id="job-1",
            sections=[
                StructuredSectionDraft(
                    section_id="s1",
                    title="Repair Procedure",
                    blocks=[
                        StructuredBlock(
                            block_id="l1",
                            block_type="bullet_list",
                            items=[StructuredListItem(content_md="Check **valve**.")],
                        ),
                        StructuredBlock(
                            block_id="n1",
                            block_type="numbered_list",
                            items=[StructuredListItem(content_md="Power off machine.")],
                        ),
                    ],
                )
            ],
        ),
        output_path,
    )

    document = Document(output_path)
    texts = [paragraph.text for paragraph in document.paragraphs]
    assert "- Check valve." in texts
    assert "1. Power off machine." in texts
    assert "**" not in "\n".join(texts)


def test_provenance_report_flattens_rich_block_evidence():
    generation = GenerationResult(
        job_id="job-1",
        sections=[
            StructuredSectionDraft(
                section_id="s1",
                title="Repair Procedure",
                blocks=[
                    StructuredBlock(
                        block_id="list-1",
                        block_type="bullet_list",
                        items=[
                            StructuredListItem(content_md="Check valve.", source_chunk_ids=["source-1-c1"]),
                            StructuredListItem(content_md="Check sensor.", reference_item_ids=["ref-1-i1"]),
                        ],
                    )
                ],
            )
        ],
    )

    report = build_provenance_report(type("Plan", (), {"job_id": "job-1", "sections": []})(), generation)

    blocks = report["sections"][0]["blocks"]
    assert blocks[0]["block_id"] == "list-1-i1"
    assert blocks[0]["text"] == "Check valve."
    assert blocks[0]["source_evidence_ids"] == ["source-1-c1"]
    assert blocks[1]["block_id"] == "list-1-i2"
    assert blocks[1]["reference_evidence_ids"] == ["ref-1-i1"]
