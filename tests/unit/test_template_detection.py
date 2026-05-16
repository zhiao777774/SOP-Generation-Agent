from pathlib import Path

from docx import Document

from backend.app.ingestion.document_loaders import load_template_docx


def test_template_detection_reads_section_titles_from_docx_tables(tmp_path: Path):
    path = tmp_path / "table-template.docx"
    document = Document()
    table = document.add_table(rows=3, cols=2)
    table.rows[0].cells[0].text = "Purpose"
    table.rows[0].cells[1].text = "Describe the maintenance objective."
    table.rows[1].cells[0].text = "Safety Requirements"
    table.rows[1].cells[1].text = "Lockout and pressure release requirements."
    table.rows[2].cells[0].text = "Repair Procedure"
    table.rows[2].cells[1].text = "Step-by-step repair content goes here."
    document.save(path)

    template = load_template_docx(str(path))

    titles = [section.title for section in template.sections]
    assert titles == ["Purpose", "Safety Requirements", "Repair Procedure"]
    assert all(title != "General SOP" for title in titles)
    assert any("DOCX tables" in warning for warning in template.warnings)


def test_template_detection_excludes_title_and_metadata_labels(tmp_path: Path):
    path = tmp_path / "metadata-template.docx"
    document = Document()
    document.add_paragraph("AXM2 Repair SOP", style="Title")
    table = document.add_table(rows=4, cols=2)
    table.rows[0].cells[0].text = "Title"
    table.rows[0].cells[1].text = "AXM2 Repair SOP"
    table.rows[1].cells[0].text = "Document No."
    table.rows[1].cells[1].text = "SOP-AXM2-001"
    table.rows[2].cells[0].text = "Safety Requirements"
    table.rows[2].cells[1].text = "Lockout and pressure release requirements."
    table.rows[3].cells[0].text = "Repair Procedure"
    table.rows[3].cells[1].text = "Step-by-step repair content goes here."
    document.save(path)

    template = load_template_docx(str(path))

    titles = [section.title for section in template.sections]
    assert titles == ["Safety Requirements", "Repair Procedure"]
    assert "Title" not in titles
    assert "AXM2 Repair SOP" not in titles
    assert any(
        block.text == "Title" and block.metadata.get("looks_like_metadata_label") == "true"
        for block in template.blocks
    )
    assert any(
        block.text == "AXM2 Repair SOP" and block.metadata.get("looks_like_document_title") == "true"
        for block in template.blocks
    )
